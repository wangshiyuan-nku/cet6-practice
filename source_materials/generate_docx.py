"""
generate_docx.py - Seed Word documents with 10 sentences each.
Creates 3 level docs + a state file tracking pool cursor and active ranks.
map_remaining.py handles dynamic refill back to 10 after deletions.
"""
import json, os, sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

INPUT_PATH = r'D:\Desktop\四六级\六级单词学习\master_vocabulary.json'
OUTPUT_DIR = r'D:\Desktop\四六级\六级单词学习\output'
STATE_PATH = os.path.join(OUTPUT_DIR, 'state.json')

SEED_SIZE = 10  # initial sentences per document

LEVELS = [
    {'name': 'Level_1_Basic', 'label': '基础级', 'start': 0, 'end': 2664,
     'desc': '常用词汇 — 基本表达与日常话题'},
    {'name': 'Level_2_Intermediate', 'label': '进阶级', 'start': 2664, 'end': 5328,
     'desc': '中频词汇 — 学术与讨论'},
    {'name': 'Level_3_Advanced', 'label': '高级', 'start': 5328, 'end': 99999,
     'desc': '低频词汇 — 深度阅读与专业领域'},
]


def sanitize_text(text):
    if not text:
        return ''
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text


def setup_styles(doc):
    en_style = doc.styles.add_style('EnglishSentence', WD_STYLE_TYPE.PARAGRAPH)
    en_style.font.name = 'Calibri'
    en_style.font.size = Pt(11)
    en_style.font.color.rgb = RGBColor(0, 0, 0)
    en_style.paragraph_format.space_after = Pt(2)
    en_style.paragraph_format.line_spacing = 1.3

    zh_style = doc.styles.add_style('ChineseTranslation', WD_STYLE_TYPE.PARAGRAPH)
    zh_style.font.name = 'SimSun'
    zh_style.font.size = Pt(9)
    zh_style.font.color.rgb = RGBColor(140, 140, 140)
    zh_style.font.italic = True
    zh_style.paragraph_format.space_before = Pt(300)
    zh_style.paragraph_format.space_after = Pt(18)
    zh_style.paragraph_format.line_spacing = 1.2

    rPr = zh_style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), 'SimSun')
    rPr.insert(0, rFonts)


def add_instruction_page(doc, level_info):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'CET-6 句子练习 — {level_info["label"]}')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(level_info['desc'])
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    instructions = [
        ('使用说明', True, 14), ('', False, 10),
        ('1. 阅读英文句子，尝试在脑中翻译成中文。', False, 10), ('', False, 6),
        ('2. 往下滚动查看中文翻译，对照检查。', False, 10), ('', False, 6),
        ('3. 看懂了 → 选中整个条目（序号+英文+空行+中文）→ Delete 删除。', False, 10), ('', False, 6),
        ('4. 没看懂 → 保留不动。', False, 10), ('', False, 6),
        ('5. 保存文档后运行 python map_remaining.py，自动标记已掌握并补满至 10 句。', False, 10), ('', False, 6),
        ('提示：蓝色 #编号 用于定位，不影响练习。', False, 9),
    ]
    for text, bold, size in instructions:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if not bold and size <= 10:
            run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_page_break()


def create_seed_document(words, level_info, output_path):
    """Create a document with exactly len(words) sentences (seed = 10)."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_instruction_page(doc, level_info)

    for w in words:
        # Word label
        word_label = doc.add_paragraph()
        word_label.paragraph_format.space_after = Pt(0)
        run = word_label.add_run(sanitize_text(f'#{w["freq_rank"]}'))
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # English sentence
        en_para = doc.add_paragraph()
        en_para.style = doc.styles['EnglishSentence']
        if en_para.runs and en_para.runs[0].text == '':
            en_para.clear()
        run = en_para.add_run(sanitize_text(w['sentence_en']))
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

        # Chinese translation
        zh_para = doc.add_paragraph()
        zh_para.style = doc.styles['ChineseTranslation']
        if zh_para.runs and zh_para.runs[0].text == '':
            zh_para.clear()
        run = zh_para.add_run(sanitize_text(w['sentence_zh']))
        run.font.size = Pt(9)
        run.font.name = 'SimSun'
        run.font.color.rgb = RGBColor(140, 140, 140)
        run.font.italic = True

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return len(words)


def build_pool():
    """Build the frequency-ordered word pool from master_vocabulary.json."""
    with open(INPUT_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return [w for w in data['words']
            if w.get('sentence_en') and w['sentence_en'] != '[暂无例句]' and w['freq_rank'] <= 8014]


def main():
    pool = build_pool()
    print(f"Total word pool: {len(pool)}")

    # Calculate level boundaries by index in pool
    n = len(pool)
    boundaries = [
        ('Level_1_Basic', 0, n // 3),
        ('Level_2_Intermediate', n // 3, 2 * n // 3),
        ('Level_3_Advanced', 2 * n // 3, n),
    ]

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    state = {'levels': {}, 'mastered_ranks': []}

    for name, start_idx, end_idx in boundaries:
        level_pool = pool[start_idx:end_idx]
        seed_words = level_pool[:SEED_SIZE]
        cursor = SEED_SIZE  # next index to pull from level_pool

        filename = f'CET6_{name}.docx'
        path = os.path.join(OUTPUT_DIR, filename)
        count = create_seed_document(seed_words, {
            'name': name, 'label': name.replace('_', ' '),
            'desc': f'{len(level_pool)} 词池 | 当前 {SEED_SIZE} 句'
        }, path)
        print(f"  {name}: {count} seed sentences -> {filename} (pool: {len(level_pool)} words)")

        state['levels'][filename] = {
            'pool_size': len(level_pool),
            'cursor': cursor,
            'active_ranks': [w['freq_rank'] for w in seed_words],
            'pool_start_idx': start_idx,
            'pool_end_idx': end_idx,
        }

    with open(STATE_PATH, 'w', encoding='utf-8') as f:
        json.dump(state, f, ensure_ascii=False, indent=2)
    print(f"\nState saved: {STATE_PATH}")
    print("Ready. Run map_remaining.py after deleting sentences to refill.")


if __name__ == '__main__':
    main()
