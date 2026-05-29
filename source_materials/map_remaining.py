"""
map_remaining.py - Dynamic 10-sentence refill with adaptive learning.

After user deletes sentences in Word:
1. Detects which ranks were deleted → marks as mastered
2. Refills document back to 10 sentences
3. Prioritizes unseen words > struggling words > mastered words (spaced review)

State tracked per word: exposure count, mastered status, last seen timestamp.
"""
import json, os, sys, io, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

MASTER_PATH = r'D:\Desktop\四六级\六级单词学习\master_vocabulary.json'
OUTPUT_DIR = r'D:\Desktop\四六级\六级单词学习\output'
STATE_PATH = os.path.join(OUTPUT_DIR, 'state.json')
INDEX_PATH = os.path.join(OUTPUT_DIR, 'word_sentence_index.json')

TARGET_SIZE = 10


def sanitize_text(text):
    if not text:
        return ''
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text


def load_json(path):
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def parse_docx(docx_path):
    """Find remaining word labels (#N) in the document."""
    doc = Document(docx_path)
    ranks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and para.runs:
            run = para.runs[0]
            if run.font.size and run.font.size.pt == 7:
                m = re.match(r'^#(\d+)', text)
                if m:
                    ranks.append(int(m.group(1)))
    return ranks


def build_vocab_lookup():
    data = load_json(MASTER_PATH)
    return {w['freq_rank']: w for w in data.get('words', [])}


def pick_next_words(level_pool, state_info, word_stats, count):
    """
    Pick `count` words from level_pool, prioritizing:
    1. Unseen words (exposure = 0) — highest priority
    2. Seen but struggling (exposure > 0, not mastered) — medium priority
    3. Mastered words (spaced repetition) — lowest priority, max 1 per refill

    Returns list of word entries from the pool.
    """
    cursor = state_info.get('cursor', 0)
    active = set(state_info.get('active_ranks', []))
    pool_size = len(level_pool)

    # Classify unseen words beyond cursor
    unseen = []
    struggling = []
    mastered_review = []

    for i, w in enumerate(level_pool):
        rank = w['freq_rank']
        if rank in active:
            continue  # already in document
        stats = word_stats.get(str(rank), {})
        exposed = stats.get('exposure_count', 0)
        mastered = stats.get('mastered', False)

        if exposed == 0:
            unseen.append((i, w))
        elif mastered:
            mastered_review.append((i, w))
        else:
            struggling.append((i, w))

    picked = []

    # Priority 1: unseen words (take from nearest to cursor first)
    unseen.sort(key=lambda x: (x[0] - cursor) % pool_size)  # closest to cursor first
    for _, w in unseen:
        if len(picked) >= count:
            break
        picked.append(w)

    # Priority 2: struggling words (user didn't delete - needs more exposure)
    if len(picked) < count:
        struggling.sort(key=lambda x: x[1]['freq_rank'])
        for _, w in struggling:
            if len(picked) >= count:
                break
            picked.append(w)

    # Priority 3: mastered but due for review (max 1 per refill)
    if len(picked) < count and mastered_review:
        mastered_review.sort(key=lambda x: (word_stats.get(str(x[1]['freq_rank']), {}).get('last_seen', '')))
        w = mastered_review[0][1]
        if w not in picked:
            picked.append(w)

    # If still not enough (edge case), take next from cursor
    for i in range(cursor, cursor + pool_size):
        idx = i % pool_size
        w = level_pool[idx]
        if w['freq_rank'] not in active and w not in picked:
            picked.append(w)
            if len(picked) >= count:
                break

    return picked[:count]


def update_document(docx_path, new_words):
    """Append new word entries to the existing document."""
    doc = Document(docx_path)

    # Remove instruction pages (find page break, start after it)
    # Find where actual word entries start (first #N label)
    # We'll just append to the end

    for w in new_words:
        # Word label
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(sanitize_text(f'#{w["freq_rank"]}'))
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # English
        p = doc.add_paragraph()
        run = p.add_run(sanitize_text(w['sentence_en']))
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

        # Chinese
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(300)
        run = p.add_run(sanitize_text(w['sentence_zh']))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(140, 140, 140)
        run.font.italic = True

    doc.save(docx_path)
    return len(new_words)


def main():
    state = load_json(STATE_PATH)
    if not state or 'levels' not in state:
        print("No state.json found. Run generate_docx.py first.")
        return

    vocab = build_vocab_lookup()
    pool = [w for w in vocab.values()
            if w.get('sentence_en') and w['sentence_en'] != '[暂无例句]' and w['freq_rank'] <= 8014]
    pool.sort(key=lambda w: w['freq_rank'])

    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    word_stats = state.get('word_stats', {})
    cross_index = load_json(INDEX_PATH).get('sentence_to_words', {})
    word_to_rank = {w['word'].lower(): w['freq_rank'] for w in vocab.values()}

    total_mastered = len(state.get('mastered_ranks', []))
    newly_mastered = 0

    for filename, level_info in state['levels'].items():
        docx_path = os.path.join(OUTPUT_DIR, filename)
        if not os.path.exists(docx_path):
            print(f"  {filename}: not found, skipping")
            continue

        present = parse_docx(docx_path)
        previous_active = set(level_info.get('active_ranks', []))
        present_set = set(present)

        # Deleted = previously active but not present now
        deleted = previous_active - present_set
        still_there = previous_active & present_set

        # Mark deleted words as mastered
        for rank in deleted:
            rk = str(rank)
            if rk not in word_stats:
                word_stats[rk] = {}
            word_stats[rk]['mastered'] = True
            word_stats[rk]['last_seen'] = now
            word_stats[rk]['exposure_count'] = word_stats[rk].get('exposure_count', 0) + 1

            # Cross-mastery: words appearing in the same sentence
            words_in_sent = cross_index.get(str(rank), [])
            for word_str in words_in_sent:
                r = word_to_rank.get(word_str.lower())
                if r:
                    rk2 = str(r)
                    if rk2 not in word_stats:
                        word_stats[rk2] = {}
                    if not word_stats[rk2].get('mastered'):
                        word_stats[rk2]['mastered'] = True
                        word_stats[rk2]['last_seen'] = now
                        word_stats[rk2]['exposure_count'] = word_stats[rk2].get('exposure_count', 0) + 1

        # Mark still-present words as struggling (exposed but not mastered)
        for rank in still_there:
            rk = str(rank)
            if rk not in word_stats:
                word_stats[rk] = {}
            word_stats[rk]['exposure_count'] = word_stats[rk].get('exposure_count', 0) + 1
            word_stats[rk]['last_seen'] = now

        newly_mastered += len(deleted)

        # Calculate how many to refill
        need = TARGET_SIZE - len(present_set)
        if need > 0:
            # Build level pool
            start_idx = level_info.get('pool_start_idx', 0)
            end_idx = level_info.get('pool_end_idx', len(pool))
            level_pool = pool[start_idx:end_idx]

            new_words = pick_next_words(level_pool, level_info, word_stats, need)
            added = update_document(docx_path, new_words)
            new_ranks = [w['freq_rank'] for w in new_words]

            # Mark new words as exposed
            for rank in new_ranks:
                rk = str(rank)
                if rk not in word_stats:
                    word_stats[rk] = {}
                word_stats[rk]['exposure_count'] = word_stats[rk].get('exposure_count', 0) + 1
                word_stats[rk]['last_seen'] = now

            # Update state
            level_info['active_ranks'] = list(present_set) + new_ranks
            level_info['cursor'] = level_info.get('cursor', 0) + len(new_words)
        else:
            level_info['active_ranks'] = list(present_set)

        unseen_count = sum(1 for r in level_info['active_ranks']
                          if word_stats.get(str(r), {}).get('exposure_count', 0) <= 1)

        print(f"  {filename}: deleted {len(deleted)} | refilled {max(0, need)} | "
              f"active {len(level_info['active_ranks'])} | {unseen_count} new words")

    # Update master state
    total_mastered += newly_mastered
    state['word_stats'] = word_stats
    state['mastered_ranks'] = sorted(list(set(
        int(k) for k, v in word_stats.items() if v.get('mastered')
    )))
    state['last_updated'] = now
    state['total_mastered'] = len(state['mastered_ranks'])

    # Summary
    total_pool = sum(l['pool_size'] for l in state['levels'].values())
    stats_summary = {
        'total_pool': total_pool,
        'mastered': len(state['mastered_ranks']),
        'unseen_remaining': total_pool - len(word_stats),
        'struggling': sum(1 for k, v in word_stats.items()
                         if v.get('exposure_count', 0) > 0 and not v.get('mastered')),
        'completion': f'{100 * len(state["mastered_ranks"]) / max(total_pool, 1):.1f}%'
    }
    state['summary'] = stats_summary

    with open(STATE_PATH, 'w', encoding='utf-8') as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

    print(f"\n{'='*50}")
    print(f"Updated: {now}")
    print(f"  Newly mastered this run: {newly_mastered}")
    print(f"  Total mastered: {stats_summary['mastered']} ({stats_summary['completion']})")
    print(f"  Words struggling (seen but not mastered): {stats_summary['struggling']}")
    print(f"  Unseen words remaining: {stats_summary['unseen_remaining']}")
    print(f"  State saved: {STATE_PATH}")


if __name__ == '__main__':
    main()
